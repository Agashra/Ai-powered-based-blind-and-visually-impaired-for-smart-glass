from flask import Flask, render_template, flash, request, session
from flask import render_template, redirect, url_for, request
#from wtforms import Form, TextField, TextAreaField, validators, StringField, SubmitField
#from werkzeug.utils import secure_filename
from PIL import Image, ImageChops,ImageStat
import yagmail

import mysql.connector
import sys, fsdk, math, ctypes, time
import datetime
app = Flask(__name__)
app.config['DEBUG']
app.config['SECRET_KEY'] = '7d441f27d441f27567d441f2b6176a'

def dmail():
    ts = time.time()
    date = datetime.datetime.fromtimestamp(ts).strftime('%Y-%m-%d')
    import mysql.connector
    conn = mysql.connector.connect(user='root', password='', host='localhost', database='1facedb')
    cursor = conn.cursor()
    cur = conn.cursor()
    cur.execute("SELECT id,Userid,Status FROM attentb where date='"+date+"' and Status='absent'")
    data = cur.fetchall()
    print(data)

    # Import docx NOT python-docx
    import docx

    # Create an instance of a word document
    doc = docx.Document()

    # Add a Title to the document
    doc.add_heading('Student Attendance', 0)

    # Table data in a form of list
    # data = ((1, 'Geek 1'),(2, 'Geek 2'),(3, 'Geek 3'))

    # Creating a table object
    table = doc.add_table(rows=1, cols=3)

    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    row[0].text = 'Id'
    row[1].text = 'RegNo'
    row[2].text = 'Status'

    # Adding data from the list to the table
    for id, RegNo, Status in data:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(id)
        row[1].text = RegNo
        row[2].text = Status

    # Now save the document to a location
    table.style = 'Colorful List'
    doc.save('record.docx')

@app.route("/")
def homepage():
    return render_template('index.html')

@app.route("/facetrain")
def facetrain():
    import LiveRecognition  as liv

    liv.att()
    del sys.modules["LiveRecognition"]
    return render_template('NewUser.html')

@app.route("/NewFace")
def NewFace():
    import LiveRecognition  as liv

    liv.att()
    del sys.modules["LiveRecognition"]
    return render_template('AdminHome.html')

@app.route("/NewStudent1", methods=['GET', 'POST'])
def NewStudent1():
     if request.method == 'POST':
          regno = request.form['regno']
          name = request.form['name']
          gender = request.form['gender']
          Age = request.form['Age']
          email = request.form['email']
          pnumber = request.form['pnumber']
          address = request.form['address']
          degg = ''
          dept =''
          year = ''
          

          conn = mysql.connector.connect(user='root', password='', host='localhost', database='1facedb')
          cursor = conn.cursor()
          cursor.execute("insert into regtb values('"+regno+"','"+name+"','"+gender+"','"+Age+"','"+email+"','"+pnumber+"','"+address +"','"+degg+"','"+dept+"','"+year+"')")
          conn.commit()
          conn.close()

     conn = mysql.connector.connect(user='root', password='', host='localhost', database='1facedb')

     cur = conn.cursor()
     cur.execute("SELECT * FROM regtb")
     data = cur.fetchall()
     return render_template('AdminHome.html', data=data)

@app.route("/searchid")
def searchid():
   # eid= request.args.get('eid')
    #session['eid']=eid



    id1 = 0
    id2 = 0

    import LiveRecognition1  as liv1
    liv1.examvales()

    #liv1.att()

    #print(ExamName)
    l1=[]

    #del sys.modules["LiveRecognition1"]




    # Fillattendances()
    n=session['n1']


    return n
@app.route("/newob")
def newob():
    import moduel_blind_object_detection


    return render_template('index.html')


@app.route("/newcur")
def newcur():
    import currency


    return render_template('index.html')



if __name__ == '__main__':
    app.run(debug=True, use_reloader=True)